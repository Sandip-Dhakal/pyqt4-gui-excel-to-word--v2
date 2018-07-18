import xlrd
import random
from docx import Document
from docx.shared import Pt

class Transformer:
    
    def __init__(self, filename='no_file', count=[10]):
        self.document = Document()
        self.answer = Document()
        self.filename = filename
        self.count = count
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        paragraph_format = style.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
        paragraph_format.line_spacing = 1.15

    def initiate(self):
        xl = xlrd.open_workbook(self.filename)
        sheetnames = xl.sheet_names()
        for idx, sheet_name in enumerate(xl.sheets()):
            total_rows = sheet_name.nrows
            total_required = self.count[idx]
            if(total_required >= total_rows):
                n = range(1,total_rows)
            else:    
                n = random.sample(range(1, total_rows), total_required)
            p = self.document.add_paragraph()
            p.add_run(sheetnames[idx]).bold = True
            pa = self.answer.add_paragraph()
            pa.add_run(sheetnames[idx]).bold = True
            for i in n:
                q = self.document.add_paragraph(style='ListNumber')
                q.add_run(sheet_name.cell(i, 1).value).bold = True
                self.document.add_paragraph('a) '+str(sheet_name.cell(i, 2).value), style='List')
                self.document.add_paragraph('b) '+str(sheet_name.cell(i, 3).value), style='List')
                self.document.add_paragraph('c) '+str(sheet_name.cell(i, 4).value), style='List')
                self.document.add_paragraph('d) '+str(sheet_name.cell(i, 5).value), style='List')
                self.document.add_paragraph( style='List')

                qa = self.answer.add_paragraph(style='ListNumber')
                qa.add_run(sheet_name.cell(i, 1).value).bold = True
                self.answer.add_paragraph('a) '+str(sheet_name.cell(i, 2).value), style='List')
                self.answer.add_paragraph('b) '+str(sheet_name.cell(i, 3).value), style='List')
                self.answer.add_paragraph('c) '+str(sheet_name.cell(i, 4).value), style='List')
                self.answer.add_paragraph('d) '+str(sheet_name.cell(i, 5).value), style='List')
                aa = self.answer.add_paragraph()
                aa.add_run("Answer: "+sheet_name.cell(i, 6).value.lower()).bold = True
                self.answer.add_paragraph( style='List')
        return self.document, self.answer