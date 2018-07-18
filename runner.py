import pandas as pd
import random
from docx import Document
from docx.shared import Pt

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

paragraph_format = style.paragraph_format
paragraph_format.space_before = 0
paragraph_format.space_after = 0
paragraph_format.line_spacing = 1.15

xl = pd.ExcelFile('Book1.xlsx')
count = 0
for sheet_name in xl.sheet_names:
    total_rows = xl.book.sheet_by_name(sheet_name).nrows
    total_required = 6
    df = pd.read_excel(xl, sheetname=sheet_name).sample(total_required)
    for i in df.index:
        document.add_paragraph(df['Question'][i], style='ListNumber')
        document.add_paragraph('a) '+df['a.'][i], style='List')
        document.add_paragraph('b) '+df['b.'][i], style='List')
        document.add_paragraph('c) '+df['c.'][i], style='List')
        document.add_paragraph('d) '+df['d.'][i], style='List')
        document.add_paragraph( style='List')
        print(df['Question'][i])
    print(total_rows)

document.save('demo.docx')